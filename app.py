# ==============================================================================
# Final Definitive app.py (With Corrected Word Downloader)
# ==============================================================================
from flask import Flask, request, render_template, jsonify, Response, send_file
import requests
import json
import os
import base64
import io
import re

# New library for PDF reading
import PyPDF2

# Import the specific exception we want to handle
from werkzeug.exceptions import RequestEntityTooLarge

from dotenv import load_dotenv
import openai
from docx import Document
from bs4 import BeautifulSoup
from docx.shared import Inches, Pt

load_dotenv()

# --- API Configuration ---
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

if not DEEPSEEK_API_KEY: raise ValueError("DEEPSEEK_API_KEY not found in .env file.")
if not OPENAI_API_KEY: raise ValueError("OPENAI_API_KEY not found in .env file.")

DEEPSEEK_BASE_URL = "https://api.deepseek.com"

try:
    openai_client = openai.OpenAI(api_key=OPENAI_API_KEY)
except Exception as e:
    raise RuntimeError(f"Failed to initialize OpenAI client. Check your OPENAI_API_KEY. Error: {e}")

app = Flask(__name__, template_folder='.', static_folder='.')

app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024 

@app.errorhandler(413)
@app.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(e):
    return jsonify(error=f"The uploaded files are too large. The server's limit is 100MB."), 413

@app.after_request
def add_no_cache_headers(response):
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response

# ### SECTION DEFINITIONS FOR DIAGRAM-BASED GENERATOR ###
SECTIONS_TO_GENERATE = {
    "Abstract": "Write an abstract of at least 10 sentences summarizing the paper's core topic, main arguments, and conclusions.",
    "Keywords": "Provide a list of relevant academic keywords.",
    "Introduction": "Write a detailed introduction of at least 10 sentences. Establish the topic's significance, state the paper's main objectives, and outline the structure.",
    "Research Gap": "This section has two paragraphs. Paragraph 1: Discuss the gap in the existing literature regarding the relationships between the core research constructs. Paragraph 2: Discuss the gap in the underlying theories that explain these constructs and their relationships.",
    "Theoretical Background and Literature Review": "Provide a deep and extensive literature review of at least 400 words. Discuss major theories, seminal works, and recent advancements. Use citations.",
    "Conceptual Framework and Hypotheses": "Based on the preceding literature review, develop and present a clear conceptual framework. Define each key construct. Formulate testable hypotheses about the relationships between these constructs as suggested by the existing literature.",
    "Synthesis of Empirical Evidence": "Write at least two detailed paragraphs, each 15 sentences long (approx. 600 words total). Review and synthesize findings from existing real empirical studies. Use numerous in-text citations. Also, generate a detailed HTML table summarizing key findings from several real, cited papers.",
    "Field Area": "This section has two paragraphs. Paragraph 1: Briefly describe the user-defined field area. Paragraph 2: Discuss the relationship between the research constructs and this specific field area. You MUST ONLY cite public government statistics or official reports (e.g., from national statistics offices, labor departments, central banks) to support your points.",
    "Discussion": "Write at least two detailed paragraphs, each 15 sentences long (approx. 600 words total). Discuss the synthesized findings from the literature review, common themes, contradictions, and significant gaps in existing knowledge. DO NOT mention the separate 'results image' in this section.",
    "Implications for Theory and Practice": "Write a detailed section of at least 10 sentences discussing the theoretical contributions and practical/managerial implications.",
    "Limitations of Existing Research and Future Directions": "Write a detailed section of at least 10 sentences discussing the limitations of the current literature and suggest specific directions for future research.",
    "Conclusion": "Write a strong, final summary of at least 10 sentences, restating the main arguments and the overall state of knowledge.",
    "References": "Provide an extensive list of all real, verifiable academic references cited throughout the paper. THIS IS A MANDATORY REQUIREMENT. Format them consistently."
}
RESULTS_INJECTION_START_SECTION = "Discussion"

# ### TABLE COLUMN DEFINITIONS FOR PAPER SUMMARY ###
TABLE_COLUMNS = [
    "Author and Date",
    "Journal",
    "Research Constructs Contributions", 
    "Research Gap",
    "Research Sample",
    "Statistical Techniques",
    "Research Results",
    "Theoretical Contributions + Practical Contributions",
    "Contributions",
    "Research Limitations",
    "Further Research Suggestions"
]


def analyze_image_with_openai(image_base64, mime_type, user_prompt):
    """Helper function to analyze a base64 encoded image."""
    if not image_base64 or not mime_type: return ""
    print(f"Analyzing image of type: {mime_type}...")
    analysis_prompt = (f"Analyze the provided image (diagram, chart, or data). Extract all key concepts, variables, relationships, and data points. Synthesize this into a detailed text description. The user's overall focus is: '{user_prompt if user_prompt else 'General analysis'}'")
    response = openai_client.chat.completions.create(model="gpt-4o", messages=[{"role": "user","content": [{"type": "text", "text": analysis_prompt},{"type": "image_url","image_url": {"url": f"data:{mime_type};base64,{image_base64}"}}]}], max_tokens=2000)
    return response.choices[0].message.content


def stream_paper_generation(initial_context, field_area_text, chosen_title, all_titles, results_analysis_text=""):
    """This generator function creates the paper section-by-section and streams progress."""
    try:
        if chosen_title:
            title_update = {"type": "final_title", "content": f"<h1>{chosen_title}</h1>"}
            yield f"data: {json.dumps(title_update)}\n\n"
        
        paper_context_so_far = initial_context + f"\n\nChosen Title: {chosen_title}"
        results_have_been_injected = False
        
        sections_to_run = SECTIONS_TO_GENERATE.copy()
        if not chosen_title:
            if "Abstract" in sections_to_run: del sections_to_run["Abstract"]
            if "Keywords" in sections_to_run: del sections_to_run["Keywords"]
        
        for section_title, section_instruction in sections_to_run.items():
            if section_title == "Field Area" and not field_area_text:
                continue

            progress_update = {"type": "progress", "message": f"Generating section: {section_title}..."}
            yield f"data: {json.dumps(progress_update)}\n\n"
            
            prompt_for_section = ""
            is_injection_point = results_analysis_text and section_title == RESULTS_INJECTION_START_SECTION
            
            if section_title == "Field Area":
                prompt_for_section = (f"Your current task is to write ONLY the 'Field Area' section. Start your response directly with `<h2>{section_title}</h2>`. The user has defined the specific field area as: **'{field_area_text}'**. The core research topic is: {initial_context}\n\nFollow this two-paragraph structure: **Paragraph 1:** Briefly introduce and describe this field area. **Paragraph 2:** Discuss how the core research constructs apply to this specific field area. You MUST ONLY cite public government statistics or official reports (e.g., from national statistics offices, labor departments, or central banks) to support your points.")
            elif is_injection_point:
                results_have_been_injected = True
                prompt_for_section = (f"You are continuing a paper. The context is: {initial_context}\n\n**CRITICAL NEW INFORMATION:** You will now integrate an analysis of a 'results image'. Here is that analysis:\n---RESULTS ANALYSIS---\n{results_analysis_text}\n---END RESULTS ANALYSIS---\n\nYour task for this '{section_title}' section is to **begin discussing this new results information.** Discuss how these results relate to the literature and framework you've already established. Integrate this new data naturally with the original instruction: {section_instruction}. Start your response directly with `<h2>{section_title}</h2>`.")
            elif results_have_been_injected:
                 prompt_for_section = (f"You are continuing a paper where you have already introduced specific results. The context is: {initial_context}\n\nThe results analysis you must continue to reference is: \n---RESULTS ANALYSIS---\n{results_analysis_text}\n---END RESULTS ANALYSIS---\n\nYour task is to write the '{section_title}' section. **Continue to weave in the implications of the previously introduced results naturally** as you fulfill this instruction: {section_instruction}. Start your response directly with `<h2>{section_title}</h2>`.")
            else:
                 prompt_for_section = (f"You are writing a paper with the context: {initial_context}\n\nYour task is to write ONLY the '{section_title}' section. **Do not mention the 'results image' yet.** Start with `<h2>{section_title}</h2>`, then fulfill this instruction: {section_instruction}.")

            headers = {"Content-Type": "application/json", "Authorization": f"Bearer {DEEPSEEK_API_KEY}"}
            payload = {"model": "deepseek-chat", "messages": [{"role": "system", "content": "You are a meticulous academic researcher writing one HTML section of a paper at a time."}, {"role": "user", "content": prompt_for_section}], "max_tokens": 4096}
            response = requests.post(f"{DEEPSEEK_BASE_URL}/chat/completions", headers=headers, data=json.dumps(payload))
            if response.status_code != 200: raise Exception(f"API Error during section '{section_title}': {response.text}")
            
            section_content = response.json()["choices"][0]["message"]["content"]
            section_update = {"type": "section", "content": section_content}
            yield f"data: {json.dumps(section_update)}\n\n"
            paper_context_so_far += f"\n\n{section_content}"

        unselected_titles = [title for title in all_titles if title != chosen_title]
        if unselected_titles:
            alt_title_html = "<h2>Alternative Suggested Titles</h2><ul>"
            for title in unselected_titles: alt_title_html += f"<li>{title}</li>"
            alt_title_html += "</ul>"
            yield f"data: {json.dumps({'type': 'section', 'content': alt_title_html})}\n\n"
        
        yield f"data: {json.dumps({'type': 'done'})}\n\n"
    except Exception as e:
        yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"


def read_uploaded_file(file_storage):
    """Reads content from an uploaded file (txt, docx, or pdf)."""
    filename = file_storage.filename
    if filename.endswith('.docx'):
        try:
            doc = Document(io.BytesIO(file_storage.read()))
            return '\n'.join([p.text for p in doc.paragraphs])
        except Exception as e:
            return f"[Error reading docx file: {e}]"
    elif filename.endswith('.txt'):
        try:
            return file_storage.read().decode('utf-8')
        except Exception as e:
            return f"[Error reading txt file: {e}]"
    elif filename.endswith('.pdf'):
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_storage.read()))
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text
        except Exception as e:
            return f"[Error reading pdf: {e}]"
    else:
        return f"[Unsupported file type: {filename}]. Please upload .txt, .docx, or .pdf."


def stream_combination_with_progress(p1_content, p2_content, p3_content, chosen_title):
    """Streams the combined paper by synthesizing it section by section based on a detailed template."""
    try:
        sections_to_synthesize = list(SECTIONS_TO_GENERATE.keys())
        total_sections = len(sections_to_synthesize)
        
        title_update = {"type": "final_title", "content": f"<h1>{chosen_title}</h1>"}
        yield f"data: {json.dumps(title_update)}\n\n"

        full_context = (f"--- START OF SOURCE PAPER 1 ---\n{p1_content}\n--- END OF SOURCE PAPER 1 ---\n\n"
                        f"--- START OF SOURCE PAPER 2 ---\n{p2_content}\n--- END OF SOURCE PAPER 2 ---\n\n"
                        f"--- START OF SOURCE PAPER 3 ---\n{p3_content}\n--- END OF SOURCE PAPER 3 ---\n")
        
        for i, section_title in enumerate(sections_to_synthesize):
            percentage = int(((i + 1) / total_sections) * 100)
            progress_update = {"type": "progress", "message": f"{percentage}% - Synthesizing section: {section_title}..."}
            yield f"data: {json.dumps(progress_update)}\n\n"
            
            synthesis_prompt = (f"You are an expert academic editor. Your task is to write ONLY the '{section_title}' section for a new paper titled '{chosen_title}'. You must synthesize the best content from the three source papers provided below. If the source papers have an explicit '{section_title}' section, synthesize them. If not, you must infer and construct this section from the most relevant content scattered throughout all three papers. You MUST preserve all original citations. NEVER refer to the source documents as 'Paper 1', etc. Start your response directly with `<h2>{section_title}</h2>`.\n\n--- SOURCE PAPERS ---\n{full_context}")
            
            if section_title == "Synthesis of Empirical Evidence":
                synthesis_prompt += ("\n   - For this specific section, you are also required to generate a detailed HTML table (`<table>`) that summarizes key findings, methodologies, or constructs from several real, cited papers mentioned in the source texts.")

            headers = {"Content-Type": "application/json", "Authorization": f"Bearer {DEEPSEEK_API_KEY}"}
            payload = { "model": "deepseek-chat", "messages": [{"role": "user", "content": synthesis_prompt}], "max_tokens": 4096 }
            
            response = requests.post(f"{DEEPSEEK_BASE_URL}/chat/completions", headers=headers, data=json.dumps(payload))
            response.raise_for_status()

            section_content = response.json()["choices"][0]["message"]["content"]
            section_update = {"type": "section", "content": section_content}
            yield f"data: {json.dumps(section_update)}\n\n"

        yield f"data: {json.dumps({'type': 'done'})}\n\n"
    except Exception as e:
        yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

# ==============================================================================
# ===== DEFINITIVE AUGMENTATION FUNCTION WITH TWO-PASS AI ANALYSIS (START) =====
# ==============================================================================
def get_headings_with_ai(base_paper_content):
    """
    Pass 1: Use AI for a lightweight task - identifying section headings.
    This is fast and avoids timeouts.
    """
    prompt = (
        "You are a document structure analyzer. Your task is to read the following paper and identify all of its section headings in the order they appear. "
        "Return a VALID JSON object with a single key 'headings', which is an array of strings. "
        "Example: {\"headings\": [\"1. Introduction\", \"2. Literature Review\", \"3. Methodology\", \"4. Conclusion\"]}\n"
        "Do NOT include the main paper title, author names, or abstract. Start with the first formal section like 'Introduction'.\n\n"
        f"--- PAPER TEXT ---\n{base_paper_content}"
    )
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {DEEPSEEK_API_KEY}"}
    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "response_format": {"type": "json_object"},
        "max_tokens": 2048
    }
    try:
        response = requests.post(f"{DEEPSEEK_BASE_URL}/chat/completions", headers=headers, data=json.dumps(payload), timeout=120)
        response.raise_for_status()
        response_data = json.loads(response.json()['choices'][0]['message']['content'])
        return response_data.get('headings', [])
    except Exception as e:
        print(f"Error in get_headings_with_ai: {e}")
        return None


def stream_paper_augmentation(base_paper_content, source_contents, source_filenames):
    """
    This definitive version uses a two-pass AI approach for robust and reliable augmentation.
    """
    try:
        # --- Step 1: Summarize sources for context ---
        yield f"data: {json.dumps({'type': 'progress', 'message': '5% - Summarizing source papers...'})}\n\n"
        source_summaries = []
        for i, content in enumerate(source_contents):
            summary_prompt = "Summarize the key arguments, findings, and definitions from this paper for later use."
            response = openai_client.chat.completions.create(model="gpt-4o", messages=[{"role": "system", "content": summary_prompt}, {"role": "user", "content": content}], max_tokens=1500)
            summary = response.choices[0].message.content
            source_summaries.append(summary)
        
        highlight_classes = ["source1", "source2", "source3"]
        full_source_context = ""
        for i, summary in enumerate(source_summaries):
            full_source_context += f"--- SUMMARY OF SOURCE PAPER {i+1} (Filename: {source_filenames[i]}, Use class: '{highlight_classes[i]}') ---\n{summary}\n\n"

        # --- Step 2 (Pass 1): Use AI to identify all section headings ---
        yield f"data: {json.dumps({'type': 'progress', 'message': '15% - AI analyzing paper structure...'})}\n\n"
        headings = get_headings_with_ai(base_paper_content)

        if not headings:
            yield f"data: {json.dumps({'type': 'error', 'message': 'The AI could not identify section headings. The document might be unstructured or in an unsupported format.'})}\n\n"
            return

        # --- Step 3 (Local Structuring): Split the paper using the AI-identified headings ---
        yield f"data: {json.dumps({'type': 'progress', 'message': '25% - Structuring paper based on AI analysis...'})}\n\n"
        paper_sections = []
        # Create a regex pattern from the headings to split the document
        # We escape headings to treat them as literal strings
        split_pattern = '|'.join(re.escape(h) for h in headings)
        content_parts = re.split(f'({split_pattern})', base_paper_content)

        # The first part is everything before the first heading
        initial_block = content_parts[0].strip()
        paper_sections.append({'section_title': 'Title and Abstract', 'section_content': initial_block})
        
        # The rest of the parts come in pairs of (heading, content)
        remaining_parts = content_parts[1:]
        for i in range(0, len(remaining_parts), 2):
            section_title = remaining_parts[i]
            section_content = remaining_parts[i+1].strip() if (i+1) < len(remaining_parts) else ""
            paper_sections.append({'section_title': section_title, 'section_content': section_content})

        # --- Step 4: Stream the initial content and then augment each section ---
        key_html = '<div class="highlight-key"><h4>Highlight Key:</h4>'
        for i, filename in enumerate(source_filenames):
            key_html += f'<p><span class="key-color {highlight_classes[i]}"></span> New content from: {filename}</p>'
        key_html += '</div>'
        yield f"data: {json.dumps({'type': 'section', 'content': key_html})}\n\n"

        total_augment_sections = len(paper_sections)
        for i, section_data in enumerate(paper_sections):
            section_title = section_data.get('section_title', 'Untitled Section')
            section_content = section_data.get('section_content', '')
            
            # Adjust progress percentage to finish at 100%
            percentage = int(30 + ((i + 1) / total_augment_sections) * 70)
            progress_update = {"type": "progress", "message": f"{percentage}% - Augmenting: '{section_title}'..."}
            yield f"data: {json.dumps(progress_update)}\n\n"

            # The first section (Title/Abstract) is usually not augmented. We just display it.
            if i == 0:
                html_section = f"<h1>{section_title}</h1>\n<p>{section_content.replace(os.linesep, '<br>')}</p>"
                yield f"data: {json.dumps({'type': 'section', 'content': f'{html_section}\n\n'})}\n\n"
                continue
            
            # Do not augment the 'References' section itself
            if section_title.strip().lower() == 'references':
                html_section = f"<h2>{section_title}</h2>\n<p>{section_content.replace(os.linesep, '<br>')}</p>"
                yield f"data: {json.dumps({'type': 'section', 'content': f'{html_section}\n\n'})}\n\n"
                continue

            augmentation_prompt = (
                "You are an expert academic editor. Augment the 'Base Section' below by inserting 1-3 RELEVANT and SUPERIOR statements from the 'Source Summaries'. DO NOT replace or summarize original content.\n\n"
                f"**Source Summaries (to get new info from):**\n{full_source_context}\n\n"
                f"**Base Section ('{section_title}') to augment:**\n---\n{section_content}\n---\n\n"
                "**CRITICAL INSTRUCTIONS:**\n"
                "1. **Highlighting (MANDATORY):** Wrap every inserted sentence and its citation in a `<span>` tag with the correct class (e.g., `<span class=\"source1\">...</span>`).\n"
                "2. **Preserve Citations:** Use the original citations from the sources accurately.\n"
                "3. **Maintain Flow:** The final text must be coherent and readable.\n"
                "4. **Output Format:** Return ONLY the full, augmented content for this section, formatted into paragraphs with `<p>` tags. DO NOT include the `<h2>` heading."
            )
            
            payload = { "model": "deepseek-chat", "messages": [{"role": "user", "content": augmentation_prompt}], "max_tokens": 4096 }
            headers = {"Content-Type": "application/json", "Authorization": f"Bearer {DEEPSEEK_API_KEY}"}
            response = requests.post(f"{DEEPSEEK_BASE_URL}/chat/completions", headers=headers, data=json.dumps(payload), timeout=120)
            response.raise_for_status()
            augmented_content = response.json()["choices"][0]["message"]["content"]
            
            html_section = f"<h2>{section_title}</h2>\n{augmented_content.strip()}"
            yield f"data: {json.dumps({'type': 'section', 'content': f'{html_section}\n\n'})}\n\n"

        # --- Step 5: REMOVED the problematic consolidated reference generation ---
        # The process will now complete successfully after the final section.

        yield f"data: {json.dumps({'type': 'done'})}\n\n"
    except Exception as e:
        yield f"data: {json.dumps({'type': 'error', 'message': f'An error occurred during augmentation: {str(e)}'})}\n\n"

# ==============================================================================
# ===== DEFINITIVE AUGMENTATION FUNCTION (END) =====
# ==============================================================================

def stream_paper_summary_table(paper_content):
    """
    Streams the generation of a research paper summary table with predetermined columns.
    """
    try:
        yield f"data: {json.dumps({'type': 'progress', 'message': '10% - Analyzing paper structure...'})}\n\n"
        
        # Start building the complete table with vertical layout
        table_html = "<table border='1' style='border-collapse: collapse; width: 100%; margin: 20px 0; font-family: Arial, sans-serif; font-size: 14px; table-layout: fixed;'>\n"
        table_html += "<style>p { margin: 6px 0; line-height: 1.8; } strong { color: #2c3e50; } em { color: #7f8c8d; }</style>\n"
        
        # Create the table header row
        table_html += "<tr style='background-color: #4a90e2; color: white; font-weight: bold;'>\n"
        table_html += "<th style='border: 1px solid #ddd; padding: 12px 8px; text-align: left; font-size: 13px; width: 200px; min-width: 200px;'>Category</th>\n"
        table_html += "<th style='border: 1px solid #ddd; padding: 12px 8px; text-align: left; font-size: 13px; width: calc(100% - 200px);'>Summary</th>\n"
        table_html += "</tr>\n"
        
        yield f"data: {json.dumps({'type': 'progress', 'message': '20% - Extracting author and date information...'})}\n\n"
        
        # Generate content for each category
        for i, column_title in enumerate(TABLE_COLUMNS):
            percentage = int(20 + ((i + 1) / len(TABLE_COLUMNS)) * 70)
            progress_update = {"type": "progress", "message": f"{percentage}% - Analyzing: {column_title}..."}
            yield f"data: {json.dumps(progress_update)}\n\n"
            
            # Create specific prompts for each column
            if column_title == "Author and Date":
                prompt = (
                    "Extract the author(s) name(s) and publication date from this research paper. "
                    "Provide this information in a clear, concise format (e.g., 'Smith & Johnson (2023)' or 'Author et al. (2023)'). "
                    "If multiple authors, use 'et al.' format if there are more than 2 authors. "
                    "Return only the author and date information, nothing else.\n\n"
                    f"--- PAPER CONTENT ---\n{paper_content}"
                )
            elif column_title == "Journal":
                prompt = (
                    "Extract the journal name from this research paper. "
                    "Provide this information in a clear, concise format (e.g., 'Journal of Economics', 'Nature'). "
                    "Return only the journal name, nothing else.\n\n"
                    f"--- PAPER CONTENT ---\n{paper_content}"
                )
            elif column_title == "Research Constructs Contributions":
                prompt = (
                    "Identify and summarize the main research constructs and their contributions in this paper. "
                    "Focus on the key variables, concepts, or theoretical constructs that the research examines. "
                    "Provide a brief summary (2-3 sentences) of what these constructs contribute to the field. "
                    "Start directly with the constructs, no introductory phrases.\n\n"
                    f"--- PAPER CONTENT ---\n{paper_content}"
                )
            elif column_title == "Research Gap":
                prompt = (
                    "Identify the research gap(s) that this study addresses. "
                    "What problem, limitation, or unanswered question in the existing literature does this research aim to solve? "
                    "Provide a concise summary (2-3 sentences) of the main research gap(s). "
                    "Start directly with the gaps, no introductory phrases.\n\n"
                    f"--- PAPER CONTENT ---\n{paper_content}"
                )
            elif column_title == "Research Sample":
                prompt = (
                    "Describe the research sample used in this study. "
                    "Include information about the participants, sample size, sampling method, and any relevant demographic information. "
                    "Provide a concise summary (2-3 sentences) of the sample characteristics. "
                    "Start directly with the sample description, no introductory phrases.\n\n"
                    f"--- PAPER CONTENT ---\n{paper_content}"
                )
            elif column_title == "Statistical Techniques":
                prompt = (
                    "Identify the statistical techniques, analytical methods, or data analysis approaches used in this research. "
                    "Include information about any statistical tests, software used, or analytical frameworks. "
                    "Provide a concise summary (2-3 sentences) of the methodological approach. "
                    "Start directly with the techniques, no introductory phrases.\n\n"
                    f"--- PAPER CONTENT ---\n{paper_content}"
                )
            elif column_title == "Research Results":
                prompt = (
                    "Summarize the main findings and results of this research study. "
                    "Focus on the key outcomes, significant findings, and what the data revealed. "
                    "Provide a concise summary (2-3 sentences) of the most important results. "
                    "Start directly with the results, no introductory phrases.\n\n"
                    f"--- PAPER CONTENT ---\n{paper_content}"
                )
            elif column_title == "Theoretical Contributions + Practical Contributions":
                prompt = (
                    "Identify both the theoretical and practical contributions of this research. "
                    "Theoretical contributions: How does this research advance theory or understanding in the field? "
                    "Practical contributions: What are the real-world applications or implications? "
                    "Provide a concise summary (2-3 sentences) covering both aspects. "
                    "Start directly with the contributions, no introductory phrases.\n\n"
                    f"--- PAPER CONTENT ---\n{paper_content}"
                )
            elif column_title == "Contributions":
                prompt = (
                    "Summarize the overall contributions of this research to the academic field. "
                    "What does this study add to existing knowledge? "
                    "Provide a concise summary (2-3 sentences) of the study's main contributions. "
                    "Start directly with the contributions, no introductory phrases.\n\n"
                    f"--- PAPER CONTENT ---\n{paper_content}"
                )
            elif column_title == "Research Limitations":
                prompt = (
                    "Identify the limitations of this research study. "
                    "What are the constraints, weaknesses, or areas where the study falls short? "
                    "Provide a concise summary (2-3 sentences) of the main limitations. "
                    "Start directly with the limitations, no introductory phrases.\n\n"
                    f"--- PAPER CONTENT ---\n{paper_content}"
                )
            elif column_title == "Further Research Suggestions":
                prompt = (
                    "Identify suggestions for future research based on this study. "
                    "What areas need further investigation? What questions remain unanswered? "
                    "Provide a concise summary (2-3 sentences) of future research directions. "
                    "Start directly with the suggestions, no introductory phrases.\n\n"
                    f"--- PAPER CONTENT ---\n{paper_content}"
                )
            
            headers = {"Content-Type": "application/json", "Authorization": f"Bearer {DEEPSEEK_API_KEY}"}
            payload = {
                "model": "deepseek-chat", 
                "messages": [
                    {"role": "system", "content": "You are an expert research analyst. Provide concise, accurate summaries based on the paper content."}, 
                    {"role": "user", "content": prompt}
                ], 
                "max_tokens": 500
            }
            
            response = requests.post(f"{DEEPSEEK_BASE_URL}/chat/completions", headers=headers, data=json.dumps(payload))
            if response.status_code != 200:
                raise Exception(f"API Error during column '{column_title}': {response.text}")
            
            column_content = response.json()["choices"][0]["message"]["content"].strip()
            
            # Improve formatting for Word documents with limited width
            # Add better structure and spacing
            if column_title == "Author and Date":
                column_content = column_content.replace('**', '<strong>')
                column_content = column_content.replace('**', '</strong>')
            elif column_title == "Journal":
                column_content = column_content.replace('**', '<strong>')
                column_content = column_content.replace('**', '</strong>')
            elif column_title == "Research Constructs Contributions":
                # Format with clear headers and bullet points
                column_content = column_content.replace('###', '<br><strong>')
                column_content = column_content.replace('**', '</strong><br>')
                column_content = column_content.replace('1.', '<br><strong>1.</strong>')
                column_content = column_content.replace('2.', '<br><strong>2.</strong>')
                column_content = column_content.replace('3.', '<br><strong>3.</strong>')
                column_content = column_content.replace('4.', '<br><strong>4.</strong>')
                column_content = column_content.replace('-', '<br>•')
                column_content = column_content.replace('**Contribution**:', '<br><em>Contribution:</em>')
            elif column_title == "Research Gap":
                column_content = column_content.replace('1.', '<br><strong>1.</strong>')
                column_content = column_content.replace('2.', '<br><strong>2.</strong>')
                column_content = column_content.replace('3.', '<br><strong>3.</strong>')
            elif column_title == "Research Sample":
                column_content = column_content.replace('**', '<strong>')
                column_content = column_content.replace('**', '</strong>')
            elif column_title == "Statistical Techniques":
                column_content = column_content.replace('1.', '<br><strong>1.</strong>')
                column_content = column_content.replace('2.', '<br><strong>2.</strong>')
                column_content = column_content.replace('3.', '<br><strong>3.</strong>')
            elif column_title == "Research Results":
                column_content = column_content.replace('**', '<strong>')
                column_content = column_content.replace('**', '</strong>')
            elif column_title == "Theoretical Contributions + Practical Contributions":
                column_content = column_content.replace('### Theoretical Contributions:', '<br><strong>Theoretical Contributions:</strong>')
                column_content = column_content.replace('### Practical Contributions:', '<br><strong>Practical Contributions:</strong>')
            elif column_title == "Research Limitations":
                column_content = column_content.replace('1.', '<br><strong>1.</strong>')
                column_content = column_content.replace('2.', '<br><strong>2.</strong>')
                column_content = column_content.replace('3.', '<br><strong>3.</strong>')
            elif column_title == "Further Research Suggestions":
                column_content = column_content.replace('1.', '<br><strong>1.</strong>')
                column_content = column_content.replace('2.', '<br><strong>2.</strong>')
                column_content = column_content.replace('3.', '<br><strong>3.</strong>')
            
            # Replace double newlines with paragraph breaks
            column_content = column_content.replace('\n\n', '</p><p>')
            # Replace single newlines with line breaks
            column_content = column_content.replace('\n', '<br>')
            # Wrap in paragraph tags for proper spacing
            column_content = f"<p>{column_content}</p>"
            
            # Add alternating row colors for better readability
            row_style = "background-color: #f9f9f9;" if i % 2 == 0 else "background-color: #ffffff;"
            table_html += f"<tr style='{row_style}'>\n"
            table_html += f"<td style='border: 1px solid #ddd; padding: 12px 8px; text-align: left; vertical-align: top; font-weight: bold; width: 200px; min-width: 200px;'>{column_title}</td>\n"
            table_html += f"<td style='border: 1px solid #ddd; padding: 12px 8px; text-align: left; vertical-align: top; line-height: 1.8; width: calc(100% - 200px); font-size: 12px;'>{column_content}</td>\n"
            table_html += "</tr>\n"
        
        # Close the table
        table_html += "</table>"
        
        # Send the complete table as one section
        yield f"data: {json.dumps({'type': 'section', 'content': table_html})}\n\n"
        
        yield f"data: {json.dumps({'type': 'done'})}\n\n"
    except Exception as e:
        yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"


@app.route("/")
def home():
    return render_template("AI.html")

@app.route('/revise_paper', methods=['POST'])
def revise_paper():
    def stream_revision(original_text):
        try:
            yield f"data: {json.dumps({'type': 'progress', 'message': '10% - Correcting grammar and highlighting changes...'})}\n\n"
            grammar_prompt = (
                "You are a meticulous proofreader. Your task is to correct the grammar of the following text. You must not change the meaning, style, or rephrase sentences. Your only job is to fix spelling, punctuation, and grammatical errors. "
                "For every change you make, you MUST wrap the corrected text in a `<span class=\"highlight\">` tag. For example, if the original is 'the dog run fast', your output must be 'the dog <span class=\"highlight\">runs</span> fast'. Do not explain your changes.\n\n"
                f"--- TEXT TO PROOFREAD ---\n{original_text}"
            )
            response = openai_client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": grammar_prompt}])
            corrected_body = response.choices[0].message.content
            
            yield f"data: {json.dumps({'type': 'progress', 'message': '50% - Finalizing grammar check...'})}\n\n"
            key_html = '<div class="highlight-key"><h4>Highlight Key:</h4><p><span class="key-color highlight"></span> Grammatical Correction</p></div><br>'
            yield f"data: {json.dumps({'type': 'section', 'content': key_html})}\n\n"
            corrected_body_html = f"<div>{corrected_body.replace(os.linesep, '<br>') if corrected_body else ''}</div>"
            yield f"data: {json.dumps({'type': 'section', 'content': corrected_body_html})}\n\n"

            yield f"data: {json.dumps({'type': 'progress', 'message': '75% - Checking for missing references...'})}\n\n"
            reference_prompt = (
                "You are a reference validation expert. You will perform the following steps on the provided text:\n"
                "1.  **Step 1:** Scan the entire text and create a list of all in-text citations, like `(Author, YYYY)` or `[1]`.\n"
                "2.  **Step 2:** Find the 'References' section at the end of the paper and create a list of all full reference entries.\n"
                "3.  **Step 3:** Compare the two lists and identify every citation from Step 1 that is missing from Step 2.\n"
                "4.  **Step 4:** Report your findings. If all citations are referenced, state that clearly. If some are missing, provide a bulleted list of the missing ones.\n"
                "**Do not provide a general guide.** Execute the steps precisely.\n\n"
                f"--- FULL TEXT TO ANALYZE ---\n{original_text}"
            )
            response = openai_client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": reference_prompt}])
            reference_check_text = response.choices[0].message.content
            reference_check_html = f"<h2>Reference Check</h2><p>{reference_check_text.replace(os.linesep, '<br>') if reference_check_text else ''}</p>"
            yield f"data: {json.dumps({'type': 'section', 'content': reference_check_html})}\n\n"
            
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"
    try:
        paper_file = request.files.get('paper')
        if not paper_file: return jsonify({"error": "No paper file uploaded."}), 400
        original_text = read_uploaded_file(paper_file)
        if original_text.startswith("[Error") or original_text.startswith("[Unsupported"):
             return jsonify({"error": original_text}), 400
        return Response(stream_revision(original_text), mimetype='text/event-stream')
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/combine_papers', methods=['POST'])
def combine_papers():
    try:
        p1_file = request.files.get('paper1'); p2_file = request.files.get('paper2'); p3_file = request.files.get('paper3')
        chosen_title = request.form.get('chosenTitle')
        if not chosen_title: return jsonify({"error": "Paper Title is a mandatory field."}), 400
        if not (p1_file and p2_file and p3_file): return jsonify({"error": "All three paper files are required."}), 400
        p1_content = read_uploaded_file(p1_file); p2_content = read_uploaded_file(p2_file); p3_content = read_uploaded_file(p3_file)
        for i, content in enumerate([p1_content, p2_content, p3_content]):
            if content.startswith("[Error") or content.startswith("[Unsupported"): return jsonify({"error": f"Could not process paper {i+1}. Reason: {content}"}), 400
        return Response(stream_combination_with_progress(p1_content, p2_content, p3_content, chosen_title), mimetype='text/event-stream')
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/augment_paper', methods=['POST'])
def augment_paper():
    try:
        base_paper_file = request.files.get('base_paper')
        source_files = [f for f in [request.files.get('source_paper_1'), request.files.get('source_paper_2'), request.files.get('source_paper_3')] if f]
        if not base_paper_file: return jsonify({"error": "Base Paper is a mandatory field."}), 400
        if not source_files: return jsonify({"error": "At least one Source Paper is required."}), 400
        base_paper_content = read_uploaded_file(base_paper_file)
        source_contents = [read_uploaded_file(f) for f in source_files]
        source_filenames = [f.filename for f in source_files]
        for i, content in enumerate([base_paper_content] + source_contents):
             if content.startswith("[Error") or content.startswith("[Unsupported"):
                 return jsonify({"error": f"Could not process paper. Reason: {content}"}), 400
        return Response(stream_paper_augmentation(base_paper_content, source_contents, source_filenames), mimetype='text/event-stream')
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ==============================================================================
# ===== DEFINITIVE, ROBUST WORD DOWNLOADER (START) =====
# ==============================================================================
@app.route('/download_word', methods=['POST'])
def download_word():
    """
    This robust function correctly parses any generated HTML, ensuring that
    what you see in the browser is what you get in the Word document.
    """
    try:
        html_content = request.get_json().get('html_content', '')
        # Wrap the content in a body tag for reliable parsing
        soup = BeautifulSoup(f"<body>{html_content}</body>", 'html.parser')
        doc = Document()
        
        # Iterate through all top-level children in the body
        for element in soup.body.children:
            if not hasattr(element, 'name'):
                # This handles plain text not wrapped in any tag
                if element.string and element.string.strip():
                    doc.add_paragraph(element.string.strip())
                continue

            if element.name == 'div' and 'highlight-key' in element.get('class', []):
                doc.add_paragraph("Highlight Key", style='Heading 3')
                for p_tag in element.find_all('p'):
                    doc.add_paragraph(p_tag.get_text(strip=True))
                doc.add_paragraph()
            elif element.name == 'h1':
                doc.add_heading(element.get_text(strip=True), level=1)
            elif element.name == 'h2':
                doc.add_heading(element.get_text(strip=True), level=2)
            elif element.name == 'p':
                if element.get_text(strip=True):
                    doc.add_paragraph(element.get_text())
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.get_text(strip=True), style='List Bullet')
            elif element.name == 'table':
                # Handle tables with proper formatting
                table_data = []
                for row in element.find_all('tr'):
                    row_data = []
                    for cell in row.find_all(['th', 'td']):
                        # Process cell content to preserve formatting
                        cell_content = ""
                        for content in cell.contents:
                            if hasattr(content, 'name'):
                                if content.name == 'p':
                                    cell_content += content.get_text() + "\n"
                                elif content.name == 'br':
                                    cell_content += "\n"
                                elif content.name == 'strong':
                                    cell_content += content.get_text()
                                elif content.name == 'em':
                                    cell_content += content.get_text()
                                else:
                                    cell_content += content.get_text()
                            else:
                                cell_content += str(content)
                        row_data.append(cell_content.strip())
                    table_data.append(row_data)
                
                if not table_data: continue
                doc_table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                doc_table.style = 'Table Grid'
                
                # Set table formatting
                for row in doc_table.rows:
                    row.height = Inches(0.3)  # Set minimum row height
                
                for i, row_data in enumerate(table_data):
                    for j, cell_text in enumerate(row_data):
                        cell = doc_table.cell(i, j)
                        cell.text = cell_text
                        # Set cell formatting for better spacing
                        paragraph = cell.paragraphs[0]
                        paragraph.space_after = Pt(6)  # Add space after paragraphs
                        paragraph.space_before = Pt(6)  # Add space before paragraphs
                        paragraph.line_spacing = 1.2  # Set line spacing
                
                doc.add_paragraph()  # Add space after table
            # This will catch any other block-level elements and add their text
            elif element.get_text(strip=True):
                 doc.add_paragraph(element.get_text(strip=True))

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name='research_paper.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return jsonify({"error": str(e)}), 500
# ==============================================================================
# ===== DEFINITIVE, ROBUST WORD DOWNLOADER (END) =====
# ==============================================================================

@app.route('/suggest_titles', methods=['POST'])
def suggest_titles():
    """Receives files, analyzes them, encodes them, and returns a full data package."""
    try:
        diagram_image = request.files.get("diagramImage")
        results_image = request.files.get("resultsImage")
        field_area = request.form.get("fieldArea", "").strip()
        guidance = request.form.get("guidance", "").strip()
        if not field_area: return jsonify({"error": "Field Area is mandatory."}), 400
        if not diagram_image: return jsonify({"error": "Primary Diagram/Model Image is mandatory."}), 400
        diagram_base64, diagram_mime = "", ""
        if diagram_image:
            diagram_mime = diagram_image.mimetype
            diagram_base64 = base64.b64encode(diagram_image.read()).decode('utf-8')
        results_base64, results_mime = "", ""
        if results_image:
            results_mime = results_image.mimetype
            results_base64 = base64.b64encode(results_image.read()).decode('utf-8')
        diagram_analysis_text = analyze_image_with_openai(diagram_base64, diagram_mime, guidance)
        results_analysis_text = analyze_image_with_openai(results_base64, results_mime, guidance)
        context_parts = []
        if guidance: context_parts.append(f"User Guidance: '{guidance}'.")
        context_parts.append(f"Field Area: '{field_area}'.")
        if diagram_analysis_text: context_parts.append(f"Diagram Analysis: '{diagram_analysis_text}'.")
        if results_analysis_text: context_parts.append(f"Results Analysis: '{results_analysis_text}'.")
        initial_context = "\n\n".join(context_parts)
        if "title" in guidance.lower():
            title_prompt = (f"Based on the following context, generate 5-10 academic titles. The user has provided specific guidance for the title: '{guidance}'. Use this as your primary instruction. Context: {initial_context}\n\nReturn ONLY a valid JSON object with a single key 'titles' which is an array of strings. Example: {{ \"titles\": [\"Title 1\", \"Title 2\"] }}")
        else:
            title_prompt = (f"Based on the following context, generate 5-10 concise, academic titles for a research paper. Context: {initial_context}\n\nReturn ONLY a valid JSON object with a single key 'titles' which is an array of strings. Example: {{ \"titles\": [\"Title 1\", \"Title 2\"] }}")
        headers = {"Content-Type": "application/json", "Authorization": f"Bearer {DEEPSEEK_API_KEY}"}
        payload = {"model": "deepseek-chat", "messages": [{"role": "system", "content": "You are a helpful assistant that only returns valid JSON objects."}, {"role": "user", "content": title_prompt}], "max_tokens": 1024}
        response = requests.post(f"{DEEPSEEK_BASE_URL}/chat/completions", headers=headers, data=json.dumps(payload))
        if response.status_code != 200: raise Exception(f"API Error suggesting titles: {response.text}")
        response_text = response.json()["choices"][0]["message"]["content"]
        match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if not match: raise Exception(f"Could not find valid JSON in title suggestion response: {response_text}")
        json_response = json.loads(match.group())
        json_response['field_area'] = field_area; json_response['guidance'] = guidance
        json_response['diagram_base64'] = diagram_base64; json_response['diagram_mime'] = diagram_mime
        json_response['results_base64'] = results_base64; json_response['results_mime'] = results_mime
        return jsonify(json_response)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/generate_from_image", methods=["POST"])
def generate_from_image():
    """Receives a JSON payload and generates the paper."""
    try:
        data = request.get_json()
        chosen_title = data.get("chosenTitle", "").strip()
        all_titles = data.get("titles", [])
        field_area = data.get("fieldArea", "").strip()
        guidance = data.get("guidance", "").strip()
        diagram_base64 = data.get("diagram_base64", ""); diagram_mime = data.get("diagram_mime", "")
        results_base64 = data.get("results_base64", ""); results_mime = data.get("results_mime", "")
        if not chosen_title: return jsonify({"error": "A title must be chosen."}), 400
        diagram_analysis_text = analyze_image_with_openai(diagram_base64, diagram_mime, guidance)
        results_analysis_text = analyze_image_with_openai(results_base64, results_mime, guidance)
        context_parts = []
        if guidance: context_parts.append(f"User Guidance: '{guidance}'.")
        context_parts.append(f"Field Area: '{field_area}'.")
        if diagram_analysis_text: context_parts.append(f"Diagram Analysis: '{diagram_analysis_text}'.")
        if results_analysis_text: context_parts.append(f"Results Analysis: '{results_analysis_text}'.")
        initial_context = "\n\n".join(context_parts)
        return Response(stream_paper_generation(initial_context, field_area, chosen_title, all_titles, results_analysis_text or ""), mimetype='text/event-stream')
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/ask", methods=["POST"])
def ask():
    """Simplified route for text-only requests."""
    try:
        topic = request.get_json().get("topic", "").strip()
        if not topic: return jsonify({"error": "Topic is required."}), 400
        initial_context = f"The paper is a literature review about the topic: '{topic}'."
        return Response(stream_paper_generation(initial_context, field_area_text="", chosen_title="", all_titles=[], results_analysis_text=""), mimetype='text/event-stream')
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/generate_summary_table', methods=['POST'])
def generate_summary_table():
    """Route for generating research paper summary tables."""
    try:
        paper_file = request.files.get('paper')
        if not paper_file:
            return jsonify({"error": "No paper file uploaded."}), 400
        
        paper_content = read_uploaded_file(paper_file)
        if paper_content.startswith("[Error") or paper_content.startswith("[Unsupported"):
            return jsonify({"error": paper_content}), 400
        
        return Response(stream_paper_summary_table(paper_content), mimetype='text/event-stream')
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
if __name__ == "__main__":
    app.run(debug=True)